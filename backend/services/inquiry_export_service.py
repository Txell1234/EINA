"""HTML/PDF/DOCX export for prospective inquiry reports."""
from __future__ import annotations

import asyncio
import html
import io
import tempfile
from pathlib import Path
from typing import Any

from sqlalchemy import select
from sqlalchemy.ext.asyncio import AsyncSession

from models.case import Case
from services.briefing_summary_service import summarize_briefing_for_report
from services.export_backends import ExportBackendError, render_pdf_from_html
from services.inquiry_report_content import build_inquiry_executive_summary
from services.report_i18n import get_report_strings, normalize_lang
from services.report_layout import (
    build_conclusions_block,
    build_cover_page,
    build_godet_pipeline_html,
    build_morph_cards_html,
    build_report_footer,
    build_scope_dashboard,
)
from services.report_markdown import format_report_line_html
from services.report_templates import get_report_css, normalize_template, probability_kpi_html


def _esc(s: Any) -> str:
    return html.escape(str(s) if s is not None else "")


def build_case_briefing_section_html(detail: dict[str, Any]) -> str:
    report = detail.get("case_briefing_report") or {}
    text = (report.get("text") or "").strip()
    if not text:
        return ""
    case_name = detail.get("case_name") or (detail.get("inquiry_scope") or {}).get("case_name") or ""
    parts = ["<section class='report-section'><h2>Briefing del cas</h2>"]
    if case_name:
        parts.append(f"<p><strong>Cas:</strong> {_esc(case_name)}</p>")
    parts.append(f"<p class='briefing-text'>{_esc(text)}</p>")
    if report.get("truncated"):
        parts.append(
            "<p class='muted'>Resum per a l'informe "
            f"(màx. {report.get('max_words', 300)} paraules). "
            "L'anàlisi Q2FS s'ha executat sobre el briefing complet "
            f"({report.get('original_word_count', '—')} paraules).</p>"
        )
    parts.append("</section>")
    return "\n".join(parts)


async def prepare_inquiry_for_export(
    db: AsyncSession,
    detail: dict[str, Any],
    *,
    lang: str | None = None,
) -> dict[str, Any]:
    """Attach condensed case briefing for report output; analysis keeps full case text."""
    scope = detail.get("inquiry_scope") or {}
    full_briefing = scope.get("case_description") or ""
    case_name = scope.get("case_name") or ""
    case_id = detail.get("case_id")
    if case_id:
        row = await db.execute(select(Case).where(Case.id == case_id))
        case = row.scalar_one_or_none()
        if case:
            full_briefing = case.description or full_briefing
            case_name = case.name or case_name

    lang_code = normalize_lang(lang)
    briefing_report = await summarize_briefing_for_report(full_briefing, lang=lang_code)
    return {
        **detail,
        "case_name": case_name,
        "case_briefing_report": briefing_report,
    }


def build_inquiry_report_html(
    detail: dict[str, Any],
    *,
    lang: str | None = None,
    template: str | None = None,
) -> str:
    """Deterministic styled HTML report from inquiry detail payload."""
    lang_code = normalize_lang(lang or detail.get("lang"))
    strings = get_report_strings(lang_code)
    artifacts = detail.get("artifacts") or {}
    report_meta = get_report_meta_from_detail(detail)
    tpl = normalize_template(template or report_meta.get("export_template"))
    q = detail.get("question", "")
    answer = detail.get("answer") or {}
    scope_audit = detail.get("scope_audit") or {}
    morph = artifacts.get("morph_bootstrap") or {}
    monitors = artifacts.get("monitor_suggestions") or {}
    steps = detail.get("steps_log") or []
    title = report_meta.get("report_title") or f"{strings.report_title} — Q2FS"
    subtitle = q if len(q) <= 200 else q[:200] + "…"
    meta_line = (
        f"Inquiry #{detail.get('id', '—')} · {detail.get('status')} · "
        f"Mode {detail.get('mode')} · Runs {detail.get('run_count', 0)}"
    )
    prob = answer.get("probability_pct")
    try:
        prob_int = int(float(prob or 0))
    except (TypeError, ValueError):
        prob_int = 0

    parts = [
        f"<!DOCTYPE html><html lang='{lang_code}'><head><meta charset='utf-8'>",
        f"<title>{_esc(title)}</title>",
        f"<style>{get_report_css(tpl, report_type='inquiry')}</style></head><body>",
        build_cover_page(
            tpl,
            title=_esc(title),
            subtitle=_esc(subtitle),
            meta=_esc(meta_line),
            probability_pct=prob_int if answer else None,
            possibility=_esc(str(answer.get("possibility") or "")),
        ),
        build_case_briefing_section_html(detail),
        build_inquiry_executive_summary({**detail, "lang": lang_code}, lang=lang_code),
        build_godet_pipeline_html(steps, template=tpl),
        build_scope_dashboard(scope_audit, template=tpl),
    ]

    if answer:
        parts.append("<section class='report-section'><h2>Resposta determinista</h2>")
        parts.append(probability_kpi_html(answer.get("probability_pct"), str(answer.get("possibility") or "")))
        if answer.get("possibility_rationale"):
            parts.append(f"<p>{format_report_line_html(answer.get('possibility_rationale'))}</p>")
        parts.append("</section>")
        parts.append(
            build_conclusions_block(
                answer.get("conclusions") or [],
                answer.get("reasoning") or [],
                template=tpl,
            )
        )

    godet_rows = morph.get("godet_preview") or []
    if godet_rows:
        parts.append(build_morph_cards_html(godet_rows, template=tpl))

    if monitors.get("suggested_monitors"):
        parts.append("<section class='report-section'><h2>Monitors de vigilància</h2><ul>")
        for m in monitors["suggested_monitors"]:
            parts.append(f"<li>{_esc(m.get('indicator'))}</li>")
        parts.append("</ul></section>")

    parts.append(
        build_report_footer(
            tpl,
            generated_note="Generat per EINA Q2FS — conclusions deterministes, traçabilitat via audit trail",
        )
    )
    parts.append("</body></html>")
    return "\n".join(parts)


def get_report_meta_from_detail(detail: dict[str, Any]) -> dict[str, Any]:
    artifacts = detail.get("artifacts") or {}
    meta = artifacts.get("report_meta")
    if not isinstance(meta, dict):
        return {
            "is_saved": False,
            "keep_forever": False,
            "archived": False,
            "report_title": "",
            "export_template": normalize_template(None),
            "saved_at": None,
            "notes": "",
        }
    return {
        "is_saved": bool(meta.get("is_saved")),
        "keep_forever": bool(meta.get("keep_forever")),
        "archived": bool(meta.get("archived")),
        "report_title": meta.get("report_title") or "",
        "export_template": normalize_template(meta.get("export_template")),
        "saved_at": meta.get("saved_at"),
        "notes": meta.get("notes") or "",
    }


async def export_inquiry_pdf(
    detail: dict[str, Any],
    *,
    output_dir: Path | None = None,
    lang: str | None = None,
    template: str | None = None,
) -> dict[str, Any]:
    html_str = build_inquiry_report_html(detail, lang=lang, template=template)
    inquiry_id = detail.get("id", "inquiry")
    out_dir = output_dir or Path("reports")
    out_dir.mkdir(parents=True, exist_ok=True)
    pdf_path = out_dir / f"inquiry_{inquiry_id}.pdf"
    try:
        await asyncio.to_thread(render_pdf_from_html, html_str, pdf_path)
        return {"ok": True, "format": "pdf", "file_path": str(pdf_path.resolve())}
    except ExportBackendError as exc:
        return {
            "ok": False,
            "format": "pdf",
            "error": str(exc),
            "fallback": "html",
            "html": html_str,
        }


def export_inquiry_pdf_bytes(
    detail: dict[str, Any],
    *,
    lang: str | None = None,
    template: str | None = None,
) -> tuple[bytes | None, str]:
    html_str = build_inquiry_report_html(detail, lang=lang, template=template)
    with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as tmp:
        tmp_path = Path(tmp.name)
    try:
        render_pdf_from_html(html_str, tmp_path)
        return tmp_path.read_bytes(), "application/pdf"
    except ExportBackendError as exc:
        return None, str(exc)
    finally:
        tmp_path.unlink(missing_ok=True)


def build_inquiry_report_docx(
    detail: dict[str, Any],
    *,
    lang: str | None = None,
    template: str | None = None,
) -> bytes:
    """Word export for Q2FS inquiry reports."""
    try:
        from docx import Document
    except ImportError as exc:
        raise ExportBackendError("python-docx", str(exc), hint="pip install 'python-docx>=1.1.0'") from exc

    lang_code = normalize_lang(lang or detail.get("lang"))
    strings = get_report_strings(lang_code)
    report_meta = get_report_meta_from_detail(detail)
    tpl = normalize_template(template or report_meta.get("export_template"))
    _ = tpl  # reserved for future styled docx blocks

    answer = detail.get("answer") or {}
    q = detail.get("question", "")
    title = report_meta.get("report_title") or f"{strings.report_title} — Q2FS"

    doc = Document()
    doc.add_heading(title, level=0)
    doc.add_paragraph(f"Inquiry #{detail.get('id', '—')} · {detail.get('status')} · Mode {detail.get('mode')}")
    doc.add_paragraph(q)

    briefing = detail.get("case_briefing_report") or {}
    briefing_text = (briefing.get("text") or "").strip()
    if briefing_text:
        doc.add_heading("Briefing del cas", level=1)
        case_name = detail.get("case_name") or (detail.get("inquiry_scope") or {}).get("case_name")
        if case_name:
            doc.add_paragraph(f"Cas: {case_name}")
        doc.add_paragraph(briefing_text)
        if briefing.get("truncated"):
            doc.add_paragraph(
                f"Resum per a l'informe (màx. {briefing.get('max_words', 300)} paraules). "
                f"L'anàlisi s'ha executat sobre el briefing complet "
                f"({briefing.get('original_word_count', '—')} paraules)."
            )

    doc.add_heading(strings.executive_summary, level=1)
    if answer:
        doc.add_paragraph(
            f"{strings.probability}: {answer.get('probability_pct')}% · "
            f"{strings.possibility}: {answer.get('possibility')}"
        )
        if answer.get("possibility_rationale"):
            doc.add_paragraph(str(answer.get("possibility_rationale")))

    conclusions = answer.get("conclusions") or []
    if conclusions:
        doc.add_heading(strings.es_conclusions, level=2)
        for item in conclusions[:12]:
            doc.add_paragraph(str(item), style="List Bullet")

    reasoning = answer.get("reasoning") or []
    if reasoning:
        doc.add_heading("Raonament traçable", level=2)
        for r in reasoning[:8]:
            if isinstance(r, dict):
                doc.add_paragraph(f"{r.get('conclusion', '')} — {r.get('because', '')}", style="List Bullet")

    scope_audit = detail.get("scope_audit") or {}
    if scope_audit:
        doc.add_heading("Scope OSINT", level=2)
        doc.add_paragraph(
            f"{scope_audit.get('queries_run', 0)} consultes · "
            f"{scope_audit.get('kept', 0)} articles conservats"
        )

    monitors = (detail.get("artifacts") or {}).get("monitor_suggestions") or {}
    monitor_list = monitors.get("suggested_monitors") or []
    if monitor_list:
        doc.add_heading("Monitors de vigilància", level=2)
        for m in monitor_list[:10]:
            if isinstance(m, dict):
                doc.add_paragraph(str(m.get("indicator")), style="List Bullet")

    doc.add_paragraph("Generat per EINA Q2FS")
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def export_inquiry_docx_bytes(
    detail: dict[str, Any],
    *,
    lang: str | None = None,
    template: str | None = None,
) -> bytes:
    return build_inquiry_report_docx(detail, lang=lang, template=template)

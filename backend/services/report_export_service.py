"""
Prospective Report Export — PDF (WeasyPrint) and DOCX (python-docx).

Loads data from SQLAlchemy models under models.prospective.
Blocking render calls run in asyncio.to_thread.
"""
from __future__ import annotations

import asyncio
import html
import io
import json
import logging
from datetime import datetime, timezone
from pathlib import Path
from typing import Any, Optional

from sqlalchemy import select
from sqlalchemy.ext.asyncio import AsyncSession
from services.export_backends import ExportBackendError, render_pdf_from_html
from services.report_content import build_executive_summary, build_variable_profiles
from services.report_enrichment import enrich_project_bundle
from services.report_i18n import get_report_strings, normalize_lang

from models.prospective import (
    MACTORObjective,
    MACTORPosture,
    MACTORResult,
    MICMACResult,
    MorphComponent,
    ProspectiveActor,
    ProspectiveProject,
    ProspectiveScenario,
    ProspectiveVariable,
)

logger = logging.getLogger(__name__)

_DOCX_HINT = "pip install 'python-docx>=1.1.0'"


def _load_docx():
    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Pt, RGBColor
    except ImportError as exc:
        raise ExportBackendError("python-docx", str(exc), hint=_DOCX_HINT) from exc
    primary_rgb = RGBColor(0x1E, 0x3A, 0x5F)
    return Document, WD_ALIGN_PARAGRAPH, Pt, RGBColor, primary_rgb


def _ensure_reports_dir(base_dir: str | Path | None = None) -> Path:
    directory = Path(base_dir or "reports")
    directory.mkdir(parents=True, exist_ok=True)
    return directory


def _utc_now_iso() -> str:
    return datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M UTC")


async def _load_project_bundle(db: AsyncSession, project_id: int) -> Optional[dict[str, Any]]:
    """Fetch project and related rows needed for exports."""
    project_r = await db.execute(select(ProspectiveProject).where(ProspectiveProject.id == project_id))
    project = project_r.scalar_one_or_none()
    if not project:
        return None

    vars_q = (
        await db.execute(
            select(ProspectiveVariable)
            .where(ProspectiveVariable.project_id == project_id)
            .order_by(ProspectiveVariable.order_index)
        )
    ).scalars().all()

    actors_q = (
        await db.execute(
            select(ProspectiveActor)
            .where(ProspectiveActor.project_id == project_id)
            .order_by(ProspectiveActor.order_index)
        )
    ).scalars().all()

    objectives_q = (
        await db.execute(
            select(MACTORObjective)
            .where(MACTORObjective.project_id == project_id)
            .order_by(MACTORObjective.order_index)
        )
    ).scalars().all()

    posture_rows = (
        await db.execute(select(MACTORPosture).where(MACTORPosture.project_id == project_id))
    ).scalars().all()

    micmac_r = await db.execute(select(MICMACResult).where(MICMACResult.project_id == project_id))
    micmac = micmac_r.scalar_one_or_none()

    mactor_r = await db.execute(select(MACTORResult).where(MACTORResult.project_id == project_id))
    mactor_result = mactor_r.scalar_one_or_none()

    morph_q = (
        await db.execute(
            select(MorphComponent)
            .where(MorphComponent.project_id == project_id)
            .order_by(MorphComponent.order_index)
        )
    ).scalars().all()

    scenarios_q = (
        await db.execute(
            select(ProspectiveScenario)
            .where(ProspectiveScenario.project_id == project_id)
            .order_by(ProspectiveScenario.id)
        )
    ).scalars().all()

    actor_codes = [a.code for a in actors_q]
    objective_codes = [o.code for o in objectives_q]
    posture_map = {(p.actor_code, p.objective_code): p.posture_value for p in posture_rows}
    postures_matrix: list[list[int]] = []
    for ac in actor_codes:
        row = []
        for oc in objective_codes:
            row.append(int(posture_map.get((ac, oc), 0)))
        postures_matrix.append(row)

    return await enrich_project_bundle(
        db,
        {
        "project": project,
        "variables": list(vars_q),
        "actors": list(actors_q),
        "objectives": list(objectives_q),
        "actor_codes": actor_codes,
        "objective_codes": objective_codes,
        "postures_matrix": postures_matrix,
        "micmac": micmac,
        "mactor_result": mactor_result,
        "components": list(morph_q),
        "scenarios": list(scenarios_q),
        },
    )


def _prepare_export_bundle(bundle: dict[str, Any], lang: str | None = None) -> dict[str, Any]:
    bundle["lang"] = normalize_lang(lang)
    bundle["strings"] = get_report_strings(bundle["lang"])
    bundle["variable_profiles"] = build_variable_profiles(bundle)
    bundle["executive_summary"] = build_executive_summary(bundle)
    return bundle


async def _attach_decision_annex(
    db: AsyncSession,
    bundle: dict[str, Any],
    *,
    include_decision_annex: bool,
) -> dict[str, Any]:
    """Opt-in decision annex; default export unchanged when flag is false."""
    bundle["decision_annex"] = None
    if not include_decision_annex:
        return bundle
    project: ProspectiveProject = bundle["project"]
    case_id = getattr(project, "case_id", None)
    if not case_id:
        return bundle
    from services.decision_annex_service import build_decision_annex, decision_annex_html

    annex = await build_decision_annex(db, case_id, project_id=project.id)
    bundle["decision_annex"] = annex
    if annex.get("has_content"):
        bundle["decision_annex_html"] = decision_annex_html(annex)
    return bundle


def _append_decision_annex_html(parts: list[str], bundle: dict[str, Any]) -> None:
    html_fragment = bundle.get("decision_annex_html")
    if html_fragment:
        parts.append(html_fragment)


def _append_decision_annex_docx(doc: Any, bundle: dict[str, Any]) -> None:
    annex = bundle.get("decision_annex")
    if not annex or not annex.get("has_content"):
        return
    doc.add_heading("Annex de decisió (intel·ligència aplicada)", level=1)
    monitor_h = annex.get("monitor_horizons") or []
    if monitor_h:
        doc.add_paragraph(f"Horitzons de monitors: {', '.join(monitor_h)}")
    for p in (annex.get("points_of_no_return") or [])[:12]:
        doc.add_paragraph(
            f"• {p.get('title', '')} — {p.get('trigger', '')} ({p.get('horizon', '')})",
            style="List Bullet",
        )
    for a in (annex.get("key_actors") or [])[:10]:
        doc.add_paragraph(
            f"{a.get('name')}: {a.get('statement_count')} declaracions, postura {a.get('avg_posture')}",
            style="List Bullet",
        )
    sig = annex.get("signal_breakdown") or {}
    if any(sig.values()):
        doc.add_paragraph(
            f"Estructural: {sig.get('structural', 0)} · Episòdic: {sig.get('episodic', 0)} · "
            f"Sense classificar: {sig.get('unknown', 0)}"
        )


def _format_sectors_table(sectors: Any) -> list[list[str]]:
    if not isinstance(sectors, list):
        return []
    rows: list[list[str]] = []
    for s in sectors:
        if isinstance(s, dict):
            rows.append(
                [
                    str(s.get("code", "")),
                    str(s.get("sector", "")),
                    str(s.get("motricitat", s.get("motricite", ""))),
                    str(s.get("dependencia", s.get("dependence", ""))),
                ]
            )
    return rows


def _posture_label(value: int) -> str:
    labels = {
        2: "+2 molt favorable",
        1: "+1 favorable",
        0: "0 neutral",
        -1: "-1 contrari",
        -2: "-2 molt contrari",
    }
    return labels.get(value, str(value))


def _append_case_briefing_html(parts: list[str], bundle: dict[str, Any]) -> None:
    case = bundle.get("case")
    if not case:
        return
    parts.append("<h1>Briefing del cas OSINT</h1>")
    parts.append(f"<p><strong>Nom:</strong> {_escape(case.name)}</p>")
    if case.description:
        parts.append(f"<p><strong>Descripció completa:</strong></p><p>{_escape(case.description)}</p>")
    prompt = bundle.get("case_prompt")
    if prompt and prompt.ai_analysis:
        parts.append("<h2>Pla de recerca (IA)</h2>")
        parts.append(f"<pre class='json'>{_escape(_json_block(prompt.ai_analysis))}</pre>")


def _append_osint_sources_html(parts: list[str], bundle: dict[str, Any]) -> None:
    articles: list = bundle.get("osint_articles") or []
    errors: list = bundle.get("osint_query_errors") or []
    parts.append("<h1>Fonts OSINT i recollida</h1>")

    if not articles and not errors:
        parts.append("<p class='muted'>Sense fonts OSINT vinculades al cas.</p>")
        return

    parts.append(
        f"<p class='muted'>{len(articles)} articles recollits · "
        f"{len(errors)} consulta(s) amb error</p>"
    )

    if articles:
        parts.append("<h2>Articles i fonts recollides</h2>")
        parts.append(
            '<table class="grid"><thead><tr>'
            "<th>Tipus consulta</th><th>Font</th><th>Títol</th><th>URL</th>"
            "<th>Data</th><th>Consulta</th>"
            "</tr></thead><tbody>"
        )
        for a in articles[:200]:
            url = a.get("url") or ""
            url_cell = f'<a href="{_escape(url)}">{_escape(url[:80])}</a>' if url else "—"
            params = _escape(_json_block(a.get("query_params") or {}))[:80]
            parts.append(
                f"<tr><td>{_escape(a.get('query_type'))}</td>"
                f"<td>{_escape(a.get('source') or '—')}</td>"
                f"<td>{_escape(a.get('title') or '—')}</td><td>{url_cell}</td>"
                f"<td>{_escape(a.get('date') or '—')}</td>"
                f"<td>{params}</td></tr>"
            )
        parts.append("</tbody></table>")
        if len(articles) > 200:
            parts.append(f"<p class='muted'>(Mostrant 200 de {len(articles)} articles)</p>")

    if errors:
        parts.append("<h2>Errors de recollida</h2>")
        parts.append(
            '<table class="grid"><thead><tr>'
            "<th>Tipus consulta</th><th>Paràmetres</th><th>Error</th><th>Estat</th>"
            "</tr></thead><tbody>"
        )
        for e in errors:
            params = _escape(_json_block(e.get("query_params") or {}))[:120]
            parts.append(
                f"<tr><td>{_escape(e.get('query_type'))}</td><td>{params}</td>"
                f"<td>{_escape(e.get('error') or '—')}</td>"
                f"<td>{_escape(e.get('result_status'))}</td></tr>"
            )
        parts.append("</tbody></table>")


def _append_extraction_html(parts: list[str], bundle: dict[str, Any]) -> None:
    statements: list = bundle.get("statements") or []
    parts.append("<h1>Extracció estructurada (declaracions)</h1>")
    if not statements:
        parts.append("<p class='muted'>Cap declaració extreta. Executa l'extracció OSINT abans de l'anàlisi.</p>")
        return
    parts.append(f"<p class='muted'>Total: {len(statements)} declaracions</p>")
    parts.append(
        '<table class="grid"><thead><tr>'
        "<th>Actor</th><th>Declaració</th><th>Postura</th><th>Tema</th>"
        "<th>Font (URL)</th><th>Data</th><th>Grounding</th><th>Estat</th>"
        "</tr></thead><tbody>"
    )
    for s in statements[:120]:
        url = s.source_url or ""
        url_cell = f'<a href="{_escape(url)}">{_escape(url[:60])}</a>' if url else "—"
        parts.append(
            f"<tr><td>{_escape(s.actor)}</td><td>{_escape(s.statement[:400])}</td>"
            f"<td>{_escape(_posture_label(int(s.posture_value or 0)))} → {_escape(s.posture_toward)}</td>"
            f"<td>{_escape(s.topic)}</td><td>{url_cell}</td>"
            f"<td>{_escape(s.source_date)}</td>"
            f"<td>{_escape(str(s.grounding_score))}</td>"
            f"<td>{_escape(s.cleanup_decision)}</td></tr>"
        )
    if len(statements) > 120:
        parts.append(f"</tbody></table><p class='muted'>(Mostrant 120 de {len(statements)})</p>")
    else:
        parts.append("</tbody></table>")


def _append_retrospective_html(parts: list[str], bundle: dict[str, Any]) -> None:
    retro = bundle.get("retrospective")
    parts.append("<h1>Anàlisi retrospectiva</h1>")
    if not retro or not retro.get("has_data"):
        parts.append(f"<p class='muted'>{_escape((retro or {}).get('message', 'Sense dades retrospectives.'))}</p>")
        return
    parts.append(
        f"<p class='muted'>{retro.get('total_statements', 0)} declaracions · "
        f"rang: {_escape(str(retro.get('date_range')))}</p>"
    )
    key_events = retro.get("key_events") or []
    if key_events:
        parts.append("<h2>Esdeveniments clau (|postura| ≥ 2)</h2>")
        parts.append('<table class="grid"><thead><tr><th>Data</th><th>Actor</th><th>Declaració</th><th>Font</th></tr></thead><tbody>')
        for ev in key_events[:30]:
            parts.append(
                f"<tr><td>{_escape(ev.get('date'))}</td><td>{_escape(ev.get('actor'))}</td>"
                f"<td>{_escape(str(ev.get('statement', ''))[:300])}</td>"
                f"<td>{_escape(ev.get('source_url'))}</td></tr>"
            )
        parts.append("</tbody></table>")
    evidence = (retro.get("micmac_evidence") or {})
    pairs = evidence.get("pairs") or []
    if pairs:
        parts.append("<h2>Evidència OSINT per relacions MIC-MAC</h2>")
        if evidence.get("interpretation"):
            parts.append(f"<p>{_escape(evidence['interpretation'])}</p>")
        parts.append('<table class="grid"><thead><tr><th>De (tema)</th><th>Cap a</th><th>N. declaracions</th><th>Confiança</th></tr></thead><tbody>')
        for p in pairs[:40]:
            parts.append(
                f"<tr><td>{_escape(p.get('from_topic'))}</td><td>{_escape(p.get('to_topic'))}</td>"
                f"<td>{p.get('n_statements')}</td><td>{p.get('confidence')}</td></tr>"
            )
        parts.append("</tbody></table>")


def _append_qual_quant_html(parts: list[str], bundle: dict[str, Any]) -> None:
    qual: list = bundle.get("qualitative_analyses") or []
    quant: list = bundle.get("quantitative_analyses") or []
    if not qual and not quant:
        return
    parts.append("<h1>Raonament qualitatiu i quantitatiu</h1>")
    if qual:
        parts.append("<h2>Anàlisi qualitativa</h2>")
        for i, q in enumerate(qual, 1):
            parts.append(f"<h3>Anàlisi {i}</h3>")
            if q.get("framework_name"):
                parts.append(f"<p><strong>Marc:</strong> {_escape(q['framework_name'])} ({_escape(q.get('framework_type'))})</p>")
            parts.append(f"<p><strong>Premissa:</strong> {_escape(q.get('premise'))}</p>")
            parts.append(f"<p><strong>Conclusions:</strong> {_escape(q.get('conclusions'))}</p>")
            if q.get("evidence"):
                parts.append(f"<pre class='json'>{_escape(_json_block(q['evidence']))}</pre>")
            parts.append(f"<p class='muted'>Confiança: {q.get('confidence_score')}</p>")
    if quant:
        parts.append("<h2>Anàlisi quantitativa (KPIs)</h2>")
        parts.append('<table class="grid"><thead><tr><th>KPI</th><th>Tipus</th><th>Valor</th><th>Unitat</th></tr></thead><tbody>')
        for q in quant:
            parts.append(
                f"<tr><td>{_escape(q.get('kpi_name'))}</td><td>{_escape(q.get('metric_type') or q.get('kpi_type'))}</td>"
                f"<td>{_escape(q.get('value'))}</td><td>{_escape(q.get('unit'))}</td></tr>"
            )
        parts.append("</tbody></table>")


def _append_actor_impact_html(parts: list[str], bundle: dict[str, Any]) -> None:
    data: dict = bundle.get("actor_impact") or {}
    parts.append("<h1>Actors afectats i escenaris</h1>")

    if not data.get("has_data"):
        parts.append(
            "<p class='muted'>Sense anàlisi d'impacte sobre actors. "
            "Executa el pipeline d'intel·ligència o POST /api/intelligence/{case_id}/actor-impact/analyze.</p>"
        )
        return

    summary = data.get("summary") or {}
    parts.append(
        f"<p><strong>Resum:</strong> {summary.get('actor_count', 0)} actors · "
        f"{summary.get('scenario_count', 0)} escenaris · "
        f"{summary.get('claim_count', 0)} conclusions justificades · "
        f"confiança global {summary.get('overall_confidence', '—')}% · "
        f"escenari més probable: {_escape(summary.get('most_likely_scenario') or '—')}</p>"
    )

    validation = data.get("validation") or {}
    if validation and not validation.get("export_ready"):
        parts.append(
            f"<p style='color:#856404;background:#fff3cd;padding:0.5em;border-radius:4px'>"
            f"<strong>Avís de traçabilitat:</strong> "
            f"{validation.get('claims_without_citation', 0)} conclusió(ns) sense citació verificable. "
            f"Revisa l'extracció OSINT abans d'emetre l'informe.</p>"
        )

    signals = data.get("osint_signals") or {}
    if signals:
        parts.append(
            "<h2>Senyals OSINT observables</h2>"
            f"<p class='muted'>{signals.get('total_statements', 0)} declaracions · "
            f"{signals.get('hostile_statements', 0)} hostils · "
            f"{signals.get('cooperative_statements', 0)} cooperatives · "
            f"{signals.get('conflict_events', 0)} esdeveniments de conflicte · "
            f"risc geo mitjà {signals.get('avg_geopolitical_risk', '—')}/100</p>"
        )

    justifications: list = data.get("scenario_justifications") or []
    scenarios_full: list = data.get("scenarios") or []
    if scenarios_full:
        parts.append("<h2>Escenaris · possibilitat i probabilitat</h2>")
        parts.append(
            '<table class="grid"><thead><tr>'
            "<th>Escenari</th><th>Possibilitat</th><th>Probabilitat</th><th>Estimada %</th><th>Raonament</th>"
            "</tr></thead><tbody>"
        )
        for s in scenarios_full:
            parts.append(
                f"<tr><td>{_escape(s.get('name'))}</td>"
                f"<td>{_escape(s.get('possibility') or 'PLAUSIBLE')}</td>"
                f"<td>{_escape(s.get('probability_label') or s.get('probability') or '—')}</td>"
                f"<td><strong>{s.get('estimated_probability_pct', '—')}%</strong></td>"
                f"<td>{_escape(str(s.get('rationale') or '')[:220])}</td></tr>"
            )
        parts.append("</tbody></table>")
    elif justifications:
        parts.append("<h2>Justificació de probabilitat d'escenaris</h2>")
        parts.append(
            '<table class="grid"><thead><tr>'
            "<th>Escenari</th><th>Base</th><th>Estimada</th><th>Ajust</th><th>Raonament</th>"
            "</tr></thead><tbody>"
        )
        for j in justifications:
            sigs = "; ".join(j.get("supporting_signals") or [])[:120]
            parts.append(
                f"<tr><td>{_escape(j.get('scenario_name'))}</td>"
                f"<td>{j.get('base_probability_pct', '—')}%</td>"
                f"<td><strong>{j.get('estimated_probability_pct', '—')}%</strong></td>"
                f"<td>{j.get('adjustment_points', 0):+d}</td>"
                f"<td>{_escape(j.get('rationale', ''))}"
                f"{(' · Senyals: ' + _escape(sigs)) if sigs else ''}</td></tr>"
            )
        parts.append("</tbody></table>")

    claims: list = data.get("claims") or []
    if claims:
        parts.append("<h2>Conclusions justificades</h2>")
        for c in claims[:15]:
            parts.append(f"<div class='claim-block' style='margin:1em 0;padding:0.75em;border-left:3px solid var(--color-primary,#1e3a5f)'>")
            parts.append(f"<p><strong>{_escape(c.get('claim', ''))}</strong></p>")
            parts.append(
                f"<p class='muted'>Confiança: {c.get('confidence', '—')}% · "
                f"Escenari: {_escape(c.get('scenario_name', ''))} · "
                f"Mètode: {_escape(c.get('method', ''))}</p>"
            )
            evidence = c.get("evidence") or []
            if evidence:
                parts.append("<ul>")
                for ev in evidence[:3]:
                    url = ev.get("source_url") or ""
                    excerpt = _escape(str(ev.get("excerpt", ""))[:180])
                    date = _escape(ev.get("source_date") or "")
                    link = f'<a href="{_escape(url)}">{_escape(url[:60])}</a>' if url else "—"
                    parts.append(f"<li>{date} · {link} · «{excerpt}»</li>")
                parts.append("</ul>")
            parts.append("</div>")

    matrix: list = data.get("impact_matrix") or []
    if matrix:
        parts.append("<h2>Matriu actor × escenari</h2>")
        parts.append(
            '<table class="grid"><thead><tr>'
            "<th>Actor</th><th>Escenari</th><th>Impacte</th><th>Etiqueta</th>"
            "<th>Confiança</th><th>Mecanisme</th>"
            "</tr></thead><tbody>"
        )
        for row in sorted(matrix, key=lambda x: x.get("impact_score", 0))[:40]:
            parts.append(
                f"<tr><td>{_escape(row.get('actor'))}</td>"
                f"<td>{_escape(row.get('scenario_name'))}</td>"
                f"<td>{row.get('impact_score', '—')}</td>"
                f"<td>{_escape(row.get('impact_label'))}</td>"
                f"<td>{row.get('confidence', '—')}%</td>"
                f"<td>{_escape(str(row.get('mechanism', ''))[:200])}</td></tr>"
            )
        parts.append("</tbody></table>")

    actors: list = data.get("actors") or []
    if actors:
        parts.append("<h2>Inventari d'actors</h2>")
        parts.append(
            '<table class="grid"><thead><tr>'
            "<th>Actor</th><th>Tipus</th><th>Declaracions</th><th>Postura mitjana</th>"
            "<th>Risc geo</th><th>Temes</th><th>Motivació</th>"
            "</tr></thead><tbody>"
        )
        for a in actors[:20]:
            topics = ", ".join(a.get("topics") or [])[:80]
            motivation = str(a.get("motivation") or "")[:280]
            parts.append(
                f"<tr><td>{_escape(a.get('name'))}</td>"
                f"<td>{_escape(a.get('type'))}</td>"
                f"<td>{a.get('statement_count', 0)}</td>"
                f"<td>{a.get('avg_posture', '—')}</td>"
                f"<td>{a.get('geo_risk_score') if a.get('geo_risk_score') is not None else '—'}</td>"
                f"<td>{_escape(topics)}</td>"
                f"<td>{_escape(motivation)}</td></tr>"
            )
        parts.append("</tbody></table>")


def _append_tavily_research_html(parts: list[str], bundle: dict[str, Any]) -> None:
    reports: list = bundle.get("tavily_research_reports") or []
    if not reports:
        return
    parts.append("<h1>Informe Tavily Research</h1>")
    for i, rep in enumerate(reports[:3], 1):
        parts.append(f"<h2>Informe {i}</h2>")
        if rep.get("created_at"):
            parts.append(f"<p class='muted'>Generat: {_escape(str(rep.get('created_at')))}</p>")
        if rep.get("source_count"):
            parts.append(f"<p class='muted'>{rep.get('source_count')} fonts consultades</p>")
        body = str(rep.get("report") or "").replace("\n", "<br/>")
        parts.append(f"<div>{body}</div>")


def _append_report_delta_html(parts: list[str], bundle: dict[str, Any]) -> None:
    delta: dict = bundle.get("report_delta") or {}
    if not delta:
        return
    parts.append("<h2>Actualització estructurada des de l'última avaluació</h2>")
    deltas = delta.get("deltas") or {}
    parts.append(
        "<p class='muted'>"
        f"Generat: {_escape(str(delta.get('generated_at') or '—'))} · "
        f"Avaluació prèvia: {_escape(str(delta.get('assessment_saved_at') or 'cap'))}</p>"
    )
    if delta.get("has_new_data"):
        parts.append("<ul>")
        labels = {
            "statements": "Declaracions extretes",
            "alert_matches": "Coincidències d'alerta",
            "osint_results": "Resultats OSINT",
            "tavily_research_reports": "Informes Tavily Research",
        }
        for key, label in labels.items():
            n = int(deltas.get(key) or 0)
            if n:
                parts.append(f"<li>+{n} {label}</li>")
        parts.append("</ul>")
    else:
        parts.append("<p class='muted'>Sense dades noves des de l'última avaluació d'impacte.</p>")


def _append_strategic_implications_html(parts: list[str], bundle: dict[str, Any]) -> None:
    """Implicacions estratègiques derivades de l'impacte sobre actors (substitueix focus financer pur)."""
    data: dict = bundle.get("actor_impact") or {}
    if not data.get("has_data"):
        return
    parts.append("<h2>Implicacions estratègiques</h2>")
    summary = data.get("summary") or {}
    exposed = summary.get("top_exposed") or []
    if exposed:
        parts.append(
            "<p><strong>Actors més exposats (impacte negatiu):</strong> "
            + ", ".join(_escape(a) for a in exposed[:5])
            + "</p>"
        )
    likely = summary.get("most_likely_scenario")
    if likely:
        parts.append(f"<p><strong>Escenari més probable segons OSINT:</strong> {_escape(likely)}</p>")


def _append_investment_html(parts: list[str], bundle: dict[str, Any]) -> None:
    inv: dict = bundle.get("investment") or {}
    parts.append("<h1>Implicacions estratègiques i context financer</h1>")
    _append_strategic_implications_html(parts, bundle)

    if not inv.get("has_data"):
        parts.append(
            "<p class='muted'>Sense recomanació d'inversió addicional. "
            "La secció principal d'impacte sobre actors és a dalt.</p>"
        )
        return

    recs: list = inv.get("recommendations") or []
    if recs:
        parts.append("<h2>Recomanacions d'inversió</h2>")
        parts.append(
            '<table class="grid"><thead><tr>'
            "<th>Tipus</th><th>Confiança %</th><th>Raonament</th><th>Data</th>"
            "</tr></thead><tbody>"
        )
        for r in recs:
            parts.append(
                f"<tr><td>{_escape(str(r.get('type', '')).upper())}</td>"
                f"<td>{r.get('confidence', '—')}</td>"
                f"<td>{_escape(str(r.get('rationale', ''))[:400])}</td>"
                f"<td>{_escape(r.get('created_at') or '—')}</td></tr>"
            )
        parts.append("</tbody></table>")

    risks: list = inv.get("risks") or []
    if risks:
        parts.append("<h2>Riscos (geopolític, mercat, operatiu)</h2>")
        parts.append(
            '<table class="grid"><thead><tr>'
            "<th>Tipus</th><th>Nivell</th><th>%</th><th>Descripció</th><th>Factors</th>"
            "</tr></thead><tbody>"
        )
        for risk in risks[:30]:
            factors = risk.get("factors") or []
            parts.append(
                f"<tr><td>{_escape(risk.get('risk_type'))}</td>"
                f"<td>{_escape(risk.get('risk_level'))}</td>"
                f"<td>{risk.get('risk_percentage', '—')}</td>"
                f"<td>{_escape(str(risk.get('description', ''))[:250])}</td>"
                f"<td>{_escape(_json_block(factors) if factors else '—')}</td></tr>"
            )
        parts.append("</tbody></table>")

    opps: list = inv.get("opportunities") or []
    if opps:
        parts.append("<h2>Oportunitats d'inversió</h2>")
        parts.append(
            '<table class="grid"><thead><tr>'
            "<th>Títol</th><th>Descripció</th><th>Confiança</th><th>Impacte</th>"
            "</tr></thead><tbody>"
        )
        for opp in opps[:25]:
            parts.append(
                f"<tr><td>{_escape(opp.get('title'))}</td>"
                f"<td>{_escape(str(opp.get('description', ''))[:300])}</td>"
                f"<td>{opp.get('confidence', '—')}</td>"
                f"<td>{_escape(opp.get('impact_level') or '—')}</td></tr>"
            )
        parts.append("</tbody></table>")

    advisor = inv.get("advisor_analysis")
    if isinstance(advisor, dict) and not advisor.get("error"):
        parts.append("<h2>Assessor d'inversions (IA)</h2>")
        if advisor.get("recommendation"):
            parts.append(
                f"<p><strong>Recomanació:</strong> {_escape(str(advisor.get('recommendation')))} "
                f"· Confiança: {advisor.get('confidence', '—')}%</p>"
            )
        if advisor.get("rationale"):
            parts.append(f"<p><strong>Raonament:</strong> {_escape(advisor.get('rationale'))}</p>")
        for key in (
            "market_analysis",
            "risk_assessment",
            "esg_analysis",
            "geopolitical_integration",
            "timing_recommendation",
            "market_comparison",
            "company_performance",
            "opportunities",
            "key_insights",
            "metrics_extracted",
        ):
            val = advisor.get(key)
            if val:
                label = key.replace("_", " ").title()
                parts.append(f"<h3>{_escape(label)}</h3>")
                parts.append(f"<pre class='json'>{_escape(_json_block(val))}</pre>")

    for ai in inv.get("ai_analyses") or []:
        parts.append(f"<h2>Anàlisi IA ({_escape(ai.get('analysis_type'))})</h2>")
        parts.append(f"<p class='muted'>Confiança: {ai.get('confidence_score', '—')}</p>")
        if ai.get("analysis_data"):
            parts.append(f"<pre class='json'>{_escape(_json_block(ai['analysis_data']))}</pre>")


def _append_executive_summary_html(parts: list[str], bundle: dict[str, Any]) -> None:
    s = bundle["strings"]
    summary = bundle.get("executive_summary") or {}
    parts.append(f"<h1 id='executive-summary'>{_escape(s.executive_summary)}</h1>")
    for block in summary.get("sections") or []:
        parts.append(f"<h2>{_escape(block.get('title', ''))}</h2>")
        for para in block.get("paragraphs") or []:
            parts.append(f"<p>{_escape(para)}</p>")
        bullets = block.get("bullets") or []
        if bullets:
            parts.append("<ul>")
            for b in bullets:
                parts.append(f"<li>{_escape(b)}</li>")
            parts.append("</ul>")
    factors = summary.get("key_factors") or []
    if factors:
        parts.append(f"<h2>{_escape(s.key_factors)}</h2>")
        parts.append(
            '<table class="grid"><thead><tr><th>Factor</th><th>Detall</th><th>Font</th></tr></thead><tbody>'
        )
        for f in factors:
            parts.append(
                f"<tr><td>{_escape(f.get('factor'))}</td>"
                f"<td>{_escape(str(f.get('detail', ''))[:400])}</td>"
                f"<td>{_escape(f.get('source'))}</td></tr>"
            )
        parts.append("</tbody></table>")


def _append_variable_profiles_html(parts: list[str], bundle: dict[str, Any]) -> None:
    s = bundle["strings"]
    profiles: list = bundle.get("variable_profiles") or []
    parts.append(f"<h1 id='variables'>{_escape(s.variables_section)}</h1>")
    if not profiles:
        parts.append(f"<p class='muted'>{_escape(s.no_data)}</p>")
        return
    for p in profiles:
        parts.append(
            "<div class='var-profile' style='margin:1em 0;padding:0.75em 1em;"
            "border:1px solid #ddd;border-left:4px solid #1e3a5f;border-radius:4px'>"
        )
        parts.append(f"<h2>{_escape(p.get('code'))} · {_escape(p.get('name'))}</h2>")
        parts.append(f"<p><strong>{_escape(s.var_type)}:</strong> {_escape(p.get('type_label'))}</p>")
        if p.get("description"):
            parts.append(f"<p><strong>{_escape(s.var_description)}:</strong> {_escape(p.get('description'))}</p>")
        if p.get("motivation"):
            parts.append(f"<p><strong>{_escape(s.var_motivation)}:</strong> {_escape(p.get('motivation'))}</p>")
        micmac_bits = []
        if p.get("sector"):
            micmac_bits.append(f"{s.var_sector}: {p.get('sector')}")
        if p.get("motricity") is not None:
            micmac_bits.append(f"{s.var_motricity}: {p.get('motricity')}")
        if p.get("dependence") is not None:
            micmac_bits.append(f"{s.var_dependence}: {p.get('dependence')}")
        if micmac_bits:
            parts.append(
                f"<p><strong>{_escape(s.var_micmac)}:</strong> {_escape('; '.join(micmac_bits))}</p>"
            )
        if p.get("osint_rationale"):
            parts.append(
                f"<p><strong>{_escape(s.var_osint_rationale)}:</strong> {_escape(p.get('osint_rationale'))}</p>"
            )
        rels = p.get("relations") or []
        if rels:
            parts.append(f"<p><strong>{_escape(s.var_relations)}:</strong></p><ul>")
            for r in rels:
                parts.append(f"<li>{_escape(r)}</li>")
            parts.append("</ul>")
        ev = p.get("evidence") or []
        if ev:
            parts.append(f"<p><strong>{_escape(s.var_evidence)}:</strong></p><ul>")
            for e in ev:
                parts.append(f"<li>{_escape(e)}</li>")
            parts.append("</ul>")
        parts.append("</div>")


def _append_executive_summary_docx(doc: Any, bundle: dict[str, Any]) -> None:
    s = bundle["strings"]
    summary = bundle.get("executive_summary") or {}
    doc.add_heading(s.executive_summary, level=1)
    for block in summary.get("sections") or []:
        doc.add_heading(str(block.get("title", "")), level=2)
        for para in block.get("paragraphs") or []:
            doc.add_paragraph(str(para))
        for b in block.get("bullets") or []:
            doc.add_paragraph(f"• {b}")
    factors = summary.get("key_factors") or []
    if factors:
        doc.add_heading(s.key_factors, level=2)
        _doc_table(
            doc,
            ["Factor", "Detall", "Font"],
            [[f.get("factor", ""), str(f.get("detail", ""))[:400], f.get("source", "")] for f in factors],
        )


def _append_variable_profiles_docx(doc: Any, bundle: dict[str, Any]) -> None:
    s = bundle["strings"]
    profiles: list = bundle.get("variable_profiles") or []
    doc.add_heading(s.variables_section, level=1)
    if not profiles:
        doc.add_paragraph(s.no_data)
        return
    for p in profiles:
        doc.add_heading(f"{p.get('code')} · {p.get('name')}", level=2)
        doc.add_paragraph(f"{s.var_type}: {p.get('type_label', '')}")
        if p.get("description"):
            doc.add_paragraph(f"{s.var_description}: {p.get('description')}")
        if p.get("motivation"):
            doc.add_paragraph(f"{s.var_motivation}: {p.get('motivation')}")
        micmac_bits = []
        if p.get("sector"):
            micmac_bits.append(f"{s.var_sector}: {p.get('sector')}")
        if p.get("motricity") is not None:
            micmac_bits.append(f"{s.var_motricity}: {p.get('motricity')}")
        if p.get("dependence") is not None:
            micmac_bits.append(f"{s.var_dependence}: {p.get('dependence')}")
        if micmac_bits:
            doc.add_paragraph(f"{s.var_micmac}: {'; '.join(micmac_bits)}")
        if p.get("osint_rationale"):
            doc.add_paragraph(f"{s.var_osint_rationale}: {p.get('osint_rationale')}")
        for r in p.get("relations") or []:
            doc.add_paragraph(f"• {r}")
        for e in p.get("evidence") or []:
            doc.add_paragraph(f"— {e}")


def _append_micmac_reasoning_html(parts: list[str], bundle: dict[str, Any]) -> None:
    suggestions = bundle.get("micmac_suggestions")
    if not suggestions or not suggestions.get("suggestions"):
        return
    parts.append("<h2>Raonament geopolític (suggeriments matriu MIC-MAC)</h2>")
    parts.append(
        f"<p class='muted'>{suggestions.get('relations_count', 0)} relacions bilaterals · "
        f"{suggestions.get('events_count', 0)} esdeveniments diplomàtics</p>"
    )
    parts.append('<table class="grid"><thead><tr><th>Fila</th><th>Columna</th><th>Valor</th><th>Raonament</th><th>Font</th></tr></thead><tbody>')
    vars_list = bundle.get("variables") or []
    for s in suggestions.get("suggestions", [])[:50]:
        row_i, col_i = s.get("row"), s.get("col")
        row_code = vars_list[row_i].code if isinstance(row_i, int) and row_i < len(vars_list) else str(row_i)
        col_code = vars_list[col_i].code if isinstance(col_i, int) and col_i < len(vars_list) else str(col_i)
        parts.append(
            f"<tr><td>{_escape(row_code)}</td><td>{_escape(col_code)}</td>"
            f"<td>{s.get('value')}</td><td>{_escape(s.get('reason'))}</td>"
            f"<td>{_escape(s.get('source'))}</td></tr>"
        )
    parts.append("</tbody></table>")


def _doc_table(doc: Any, headers: list[str], rows: list[list[str]]) -> None:
    if not rows:
        doc.add_paragraph("Sense dades.")
        return
    table = doc.add_table(rows=1, cols=len(headers))
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)[:500]


def _append_case_briefing_docx(doc: Any, bundle: dict[str, Any]) -> None:
    case = bundle.get("case")
    if not case:
        return
    doc.add_heading("Briefing del cas OSINT", level=1)
    doc.add_paragraph(f"Nom: {case.name}")
    if case.description:
        doc.add_paragraph("Descripció completa:")
        doc.add_paragraph(case.description)
    prompt = bundle.get("case_prompt")
    if prompt and prompt.ai_analysis:
        doc.add_heading("Pla de recerca (IA)", level=2)
        doc.add_paragraph(_json_block(prompt.ai_analysis))


def _append_osint_sources_docx(doc: Any, bundle: dict[str, Any]) -> None:
    articles: list = bundle.get("osint_articles") or []
    errors: list = bundle.get("osint_query_errors") or []
    doc.add_heading("Fonts OSINT i recollida", level=1)
    if not articles and not errors:
        doc.add_paragraph("Sense fonts OSINT vinculades al cas.")
        return
    doc.add_paragraph(f"{len(articles)} articles recollits · {len(errors)} consulta(s) amb error")
    if articles:
        doc.add_heading("Articles i fonts recollides", level=2)
        rows = [
            [
                str(a.get("query_type", "")),
                str(a.get("source") or "—"),
                str(a.get("title") or "—")[:120],
                str(a.get("url") or "—")[:120],
                str(a.get("date") or "—"),
            ]
            for a in articles[:150]
        ]
        _doc_table(doc, ["Tipus", "Font", "Títol", "URL", "Data"], rows)
        if len(articles) > 150:
            doc.add_paragraph(f"(Mostrant 150 de {len(articles)} articles)")
    if errors:
        doc.add_heading("Errors de recollida", level=2)
        err_rows = [
            [
                str(e.get("query_type", "")),
                _json_block(e.get("query_params") or {})[:80],
                str(e.get("error") or "—")[:200],
                str(e.get("result_status", "")),
            ]
            for e in errors
        ]
        _doc_table(doc, ["Tipus", "Paràmetres", "Error", "Estat"], err_rows)


def _append_extraction_docx(doc: Any, bundle: dict[str, Any]) -> None:
    statements: list = bundle.get("statements") or []
    doc.add_heading("Extracció estructurada (declaracions)", level=1)
    if not statements:
        doc.add_paragraph("Cap declaració extreta.")
        return
    doc.add_paragraph(f"Total: {len(statements)} declaracions")
    rows = [
        [
            s.actor,
            (s.statement or "")[:200],
            f"{_posture_label(int(s.posture_value or 0))} → {s.posture_toward or ''}",
            s.topic or "",
            (s.source_url or "—")[:80],
            s.source_date or "",
            str(s.grounding_score or ""),
            s.cleanup_decision or "",
        ]
        for s in statements[:80]
    ]
    _doc_table(
        doc,
        ["Actor", "Declaració", "Postura", "Tema", "URL", "Data", "Grounding", "Estat"],
        rows,
    )
    if len(statements) > 80:
        doc.add_paragraph(f"(Mostrant 80 de {len(statements)})")


def _append_retrospective_docx(doc: Any, bundle: dict[str, Any]) -> None:
    retro = bundle.get("retrospective")
    doc.add_heading("Anàlisi retrospectiva", level=1)
    if not retro or not retro.get("has_data"):
        doc.add_paragraph((retro or {}).get("message", "Sense dades retrospectives."))
        return
    doc.add_paragraph(
        f"{retro.get('total_statements', 0)} declaracions · Rang: {retro.get('date_range')}"
    )
    for ev in (retro.get("key_events") or [])[:20]:
        doc.add_paragraph(
            f"• [{ev.get('date', '?')}] {ev.get('actor', '')}: {(ev.get('statement') or '')[:180]} "
            f"(Font: {ev.get('source_url', '—')})"
        )
    evidence = retro.get("micmac_evidence") or {}
    if evidence.get("interpretation"):
        doc.add_paragraph(evidence["interpretation"])
    pairs = evidence.get("pairs") or []
    if pairs:
        doc.add_heading("Evidència OSINT per MIC-MAC", level=2)
        _doc_table(
            doc,
            ["De (tema)", "Cap a", "N. declaracions", "Confiança"],
            [[p.get("from_topic", ""), p.get("to_topic", ""), str(p.get("n_statements", "")), str(p.get("confidence", ""))] for p in pairs[:30]],
        )


def _append_qual_quant_docx(doc: Any, bundle: dict[str, Any]) -> None:
    qual: list = bundle.get("qualitative_analyses") or []
    quant: list = bundle.get("quantitative_analyses") or []
    if not qual and not quant:
        return
    doc.add_heading("Raonament qualitatiu i quantitatiu", level=1)
    for i, q in enumerate(qual, 1):
        doc.add_heading(f"Anàlisi qualitativa {i}", level=2)
        if q.get("framework_name"):
            doc.add_paragraph(f"Marc: {q['framework_name']} ({q.get('framework_type')})")
        doc.add_paragraph(f"Premissa: {q.get('premise', '')}")
        doc.add_paragraph(f"Conclusions: {q.get('conclusions', '')}")
        if q.get("evidence"):
            doc.add_paragraph(_json_block(q["evidence"]))
    if quant:
        doc.add_heading("Anàlisi quantitativa (KPIs)", level=2)
        _doc_table(
            doc,
            ["KPI", "Tipus", "Valor", "Unitat"],
            [[q.get("kpi_name", ""), q.get("metric_type") or q.get("kpi_type", ""), q.get("value", ""), q.get("unit") or ""] for q in quant],
        )


def _append_actor_impact_docx(doc: Any, bundle: dict[str, Any]) -> None:
    data: dict = bundle.get("actor_impact") or {}
    doc.add_heading("Actors afectats i escenaris", level=1)
    if not data.get("has_data"):
        doc.add_paragraph(
            "Sense anàlisi d'impacte sobre actors. "
            "Executa el pipeline d'intel·ligència abans d'exportar."
        )
        return

    summary = data.get("summary") or {}
    doc.add_paragraph(
        f"Actors: {summary.get('actor_count', 0)} · Escenaris: {summary.get('scenario_count', 0)} · "
        f"Conclusions: {summary.get('claim_count', 0)} · Confiança: {summary.get('overall_confidence', '—')}% · "
        f"Escenari més probable: {summary.get('most_likely_scenario') or '—'}"
    )

    validation = data.get("validation") or {}
    if validation and not validation.get("export_ready"):
        doc.add_paragraph(
            f"AVÍS: {validation.get('claims_without_citation', 0)} conclusions sense citació verificable."
        )

    signals = data.get("osint_signals") or {}
    if signals:
        doc.add_heading("Senyals OSINT observables", level=2)
        doc.add_paragraph(
            f"{signals.get('total_statements', 0)} declaracions · "
            f"{signals.get('hostile_statements', 0)} hostils · "
            f"{signals.get('cooperative_statements', 0)} cooperatives · "
            f"{signals.get('conflict_events', 0)} esdeveniments de conflicte · "
            f"risc geo mitjà {signals.get('avg_geopolitical_risk', '—')}/100"
        )

    justifications: list = data.get("scenario_justifications") or []
    if justifications:
        doc.add_heading("Justificació de probabilitat d'escenaris", level=2)
        _doc_table(
            doc,
            ["Escenari", "Base %", "Estimada %", "Ajust", "Raonament"],
            [
                [
                    str(j.get("scenario_name", "")),
                    str(j.get("base_probability_pct", "")),
                    str(j.get("estimated_probability_pct", "")),
                    f"{j.get('adjustment_points', 0):+d}",
                    str(j.get("rationale", ""))[:150],
                ]
                for j in justifications
            ],
        )

    claims: list = data.get("claims") or []
    if claims:
        doc.add_heading("Conclusions justificades", level=2)
        for c in claims[:12]:
            doc.add_paragraph(c.get("claim", ""), style="List Bullet")
            doc.add_paragraph(
                f"Confiança {c.get('confidence', '—')}% · Escenari: {c.get('scenario_name', '')} · "
                f"Mètode: {c.get('method', '')}"
            )
            for ev in (c.get("evidence") or [])[:2]:
                doc.add_paragraph(
                    f"  Font: {ev.get('source_url') or '—'} ({ev.get('source_date') or '—'}) "
                    f"«{(ev.get('excerpt') or '')[:120]}»"
                )

    matrix: list = data.get("impact_matrix") or []
    if matrix:
        doc.add_heading("Matriu actor × escenari", level=2)
        _doc_table(
            doc,
            ["Actor", "Escenari", "Impacte", "Confiança", "Mecanisme"],
            [
                [
                    str(row.get("actor", "")),
                    str(row.get("scenario_name", "")),
                    str(row.get("impact_score", "")),
                    f"{row.get('confidence', '')}%",
                    str(row.get("mechanism", ""))[:120],
                ]
                for row in sorted(matrix, key=lambda x: x.get("impact_score", 0))[:30]
            ],
        )


def _append_investment_docx(doc: Any, bundle: dict[str, Any]) -> None:
    inv: dict = bundle.get("investment") or {}
    doc.add_heading("Implicacions estratègiques i context financer", level=1)
    if not inv.get("has_data") and not (bundle.get("actor_impact") or {}).get("has_data"):
        doc.add_paragraph("Sense dades estratègiques addicionals.")
        return
    if (bundle.get("actor_impact") or {}).get("has_data"):
        summary = (bundle["actor_impact"].get("summary") or {})
        exposed = summary.get("top_exposed") or []
        if exposed:
            doc.add_paragraph(f"Actors més exposats: {', '.join(exposed[:5])}")
    if not inv.get("has_data"):
        return

    recs: list = inv.get("recommendations") or []
    if recs:
        doc.add_heading("Recomanacions d'inversió", level=2)
        _doc_table(
            doc,
            ["Tipus", "Confiança %", "Raonament", "Data"],
            [
                [
                    str(r.get("type", "")).upper(),
                    str(r.get("confidence", "")),
                    str(r.get("rationale", ""))[:200],
                    str(r.get("created_at") or ""),
                ]
                for r in recs
            ],
        )

    risks: list = inv.get("risks") or []
    if risks:
        doc.add_heading("Riscos", level=2)
        _doc_table(
            doc,
            ["Tipus", "Nivell", "%", "Descripció"],
            [
                [
                    str(risk.get("risk_type", "")),
                    str(risk.get("risk_level", "")),
                    str(risk.get("risk_percentage", "")),
                    str(risk.get("description", ""))[:150],
                ]
                for risk in risks[:25]
            ],
        )

    opps: list = inv.get("opportunities") or []
    if opps:
        doc.add_heading("Oportunitats", level=2)
        _doc_table(
            doc,
            ["Títol", "Descripció", "Confiança", "Impacte"],
            [
                [
                    str(opp.get("title", "")),
                    str(opp.get("description", ""))[:150],
                    str(opp.get("confidence", "")),
                    str(opp.get("impact_level") or ""),
                ]
                for opp in opps[:20]
            ],
        )

    advisor = inv.get("advisor_analysis")
    if isinstance(advisor, dict) and not advisor.get("error"):
        doc.add_heading("Assessor d'inversions (IA)", level=2)
        if advisor.get("recommendation"):
            doc.add_paragraph(
                f"Recomanació: {advisor.get('recommendation')} · "
                f"Confiança: {advisor.get('confidence', '—')}%"
            )
        if advisor.get("rationale"):
            doc.add_paragraph(f"Raonament: {advisor.get('rationale')}")
        for key in ("risk_assessment", "market_analysis", "esg_analysis", "opportunities"):
            if advisor.get(key):
                doc.add_heading(key.replace("_", " ").title(), level=3)
                doc.add_paragraph(_json_block(advisor[key])[:3000])

    for ai in inv.get("ai_analyses") or []:
        doc.add_heading(f"Anàlisi IA ({ai.get('analysis_type')})", level=2)
        if ai.get("analysis_data"):
            doc.add_paragraph(_json_block(ai["analysis_data"])[:4000])


def _append_micmac_reasoning_docx(doc: Any, bundle: dict[str, Any]) -> None:
    suggestions = bundle.get("micmac_suggestions")
    if not suggestions or not suggestions.get("suggestions"):
        return
    doc.add_heading("Raonament geopolític (suggeriments MIC-MAC)", level=2)
    vars_list = bundle.get("variables") or []
    rows = []
    for s in suggestions.get("suggestions", [])[:40]:
        row_i, col_i = s.get("row"), s.get("col")
        row_code = vars_list[row_i].code if isinstance(row_i, int) and row_i < len(vars_list) else str(row_i)
        col_code = vars_list[col_i].code if isinstance(col_i, int) and col_i < len(vars_list) else str(col_i)
        rows.append([row_code, col_code, str(s.get("value", "")), str(s.get("reason", "")), str(s.get("source", ""))])
    _doc_table(doc, ["Fila", "Columna", "Valor", "Raonament", "Font"], rows)


def _json_block(data: Any) -> str:
    try:
        return json.dumps(data, indent=2, ensure_ascii=False)
    except TypeError:
        return str(data)


def _escape(s: Optional[str]) -> str:
    if s is None:
        return ""
    return html.escape(str(s))


def _html_report(bundle: dict[str, Any]) -> str:
    project: ProspectiveProject = bundle["project"]
    vars_list: list = bundle["variables"]
    actors_list: list = bundle["actors"]
    objectives_list: list = bundle["objectives"]
    micmac: Optional[MICMACResult] = bundle["micmac"]
    mactor_result: Optional[MACTORResult] = bundle["mactor_result"]
    components: list = bundle["components"]
    scenarios_list: list = bundle["scenarios"]
    actor_codes = bundle["actor_codes"]
    objective_codes = bundle["objective_codes"]
    postures_matrix = bundle["postures_matrix"]

    css = """
    @page { size: A4; margin: 18mm 16mm; }
    body { font-family: "DejaVu Sans", Helvetica, Arial, sans-serif; font-size: 10pt; color: #222; }
    h1 { color: #1e3a5f; font-size: 18pt; margin-top: 0; border-bottom: 2px solid #ff6b35; padding-bottom: 6px; }
    h2 { color: #1e3a5f; font-size: 12pt; margin-top: 16px; }
    .muted { color: #555; font-size: 9pt; }
    table.grid { border-collapse: collapse; width: 100%; margin: 8px 0 12px 0; }
    table.grid th, table.grid td { border: 1px solid #ccc; padding: 4px 6px; text-align: left; vertical-align: top; }
    table.grid th { background: #f3f6f9; font-weight: 600; }
    pre.json { white-space: pre-wrap; font-size: 8.5pt; background: #f8f9fa; padding: 8px; border-radius: 4px; }
    .cover-title { font-size: 26pt; color: #1e3a5f; margin-bottom: 4px; }
    .cover-sub { font-size: 14pt; color: #333; margin-bottom: 18px; }
    """

    s = bundle["strings"]
    parts: list[str] = []
    parts.append(f"<style>{css}</style>")
    parts.append(f'<div class="cover-title">{_escape(s.report_title)}</div>')
    parts.append(f"<div class='cover-sub'>{_escape(project.title)}</div>")
    parts.append(
        f"<p class='muted'>{_escape(s.report_subtitle)}</p>"
        f"<p class='muted'>{_escape(s.project_id)} {project.id} · "
        f"{_escape(s.generated)} {_escape(_utc_now_iso())}</p>"
    )

    _append_executive_summary_html(parts, bundle)

    parts.append(f"<h1>{_escape(s.project_section)}</h1>")
    parts.append(f"<p><strong>{_escape(s.hypothesis)}</strong></p>")
    parts.append(f"<p>{_escape(project.hypothesis)}</p>")
    parts.append(f"<p><strong>{_escape(s.context)}</strong></p>")
    parts.append(f"<p>{_escape(project.context)}</p>")

    _append_case_briefing_html(parts, bundle)
    _append_osint_sources_html(parts, bundle)
    _append_tavily_research_html(parts, bundle)
    _append_extraction_html(parts, bundle)
    _append_retrospective_html(parts, bundle)
    _append_actor_impact_html(parts, bundle)
    _append_report_delta_html(parts, bundle)
    _append_investment_html(parts, bundle)

    _append_variable_profiles_html(parts, bundle)

    parts.append("<h1>Resultats MIC-MAC</h1>")
    if micmac:
        vb = micmac.vb_index
        vr = micmac.vr_index
        meta = (
            f"<p class='muted'>Calculat: {_escape(micmac.calculated_at.isoformat()) if micmac.calculated_at else '—'}"
            f" · VB index: {_escape(str(vb) if vb is not None else '—')}"
            f" · VR index: {_escape(str(vr) if vr is not None else '—')} </p>"
        )
        parts.append(meta)
        sector_rows = _format_sectors_table(micmac.sectors)
        if sector_rows:
            parts.append("<h2>Classificació per sectors (Godet)</h2>")
            parts.append('<table class="grid"><thead><tr><th>Codi</th><th>Sector</th><th>Motricitat</th><th>Dependència</th></tr></thead><tbody>')
            for row in sector_rows:
                parts.append(f"<tr><td>{_escape(row[0])}</td><td>{_escape(row[1])}</td><td>{_escape(row[2])}</td><td>{_escape(row[3])}</td></tr>")
            parts.append("</tbody></table>")
        _append_micmac_reasoning_html(parts, bundle)
        for label, blob in (
            ("Sectors", micmac.sectors),
            ("Motricitat directa", micmac.motricite_direct),
            ("Dependència directa", micmac.dependence_direct),
            ("Matriu directa", micmac.matrix_direct),
            ("Matriu indirecta", micmac.matrix_indirect),
        ):
            if blob is not None:
                parts.append(f"<h2>{_escape(label)}</h2><pre class='json'>{_escape(_json_block(blob))}</pre>")
    else:
        parts.append("<p class='muted'>Encara sense càlcul MIC-MAC desat.</p>")

    parts.append("<h1>MACTOR · Actors i objectius</h1>")
    suggested_actors = bundle.get("suggested_actors") or []
    if suggested_actors:
        parts.append("<h2>Actors suggerits per l'extracció OSINT</h2>")
        parts.append('<table class="grid"><thead><tr><th>Codi</th><th>Nom</th><th>Força</th><th>Declaracions</th></tr></thead><tbody>')
        for sa in suggested_actors[:15]:
            parts.append(
                f"<tr><td>{_escape(sa.get('code'))}</td><td>{_escape(sa.get('name'))}</td>"
                f"<td>{_escape(str(sa.get('force')))}</td><td>{sa.get('statement_count', sa.get('n_statements', ''))}</td></tr>"
            )
        parts.append("</tbody></table>")
    if actors_list:
        parts.append('<table class="grid"><thead><tr><th>Codi</th><th>Nom</th><th>Força</th><th>Fins estratègics</th></tr></thead><tbody>')
        for a in actors_list:
            fins = ", ".join(a.strategic_goals or [])
            parts.append(
                f"<tr><td>{_escape(a.code)}</td><td>{_escape(a.name)}</td>"
                f"<td>{_escape(str(a.force_score))}</td><td>{_escape(fins)}</td></tr>"
            )
        parts.append("</tbody></table>")
    if objectives_list:
        parts.append("<h2>Objectius</h2>")
        parts.append('<table class="grid"><thead><tr><th>Codi</th><th>Nom</th></tr></thead><tbody>')
        for o in objectives_list:
            parts.append(f"<tr><td>{_escape(o.code)}</td><td>{_escape(o.name)}</td></tr>")
        parts.append("</tbody></table>")

    if actor_codes and objective_codes and postures_matrix:
        parts.append("<h2>Matriu de postures</h2>")
        parts.append('<table class="grid"><thead><tr><th>Actor \\ Objectiu</th>')
        parts.extend(f"<th>{_escape(oc)}</th>" for oc in objective_codes)
        parts.append("</tr></thead><tbody>")
        for i, ac in enumerate(actor_codes):
            parts.append(f"<tr><th>{_escape(ac)}</th>")
            row = postures_matrix[i] if i < len(postures_matrix) else []
            for j in range(len(objective_codes)):
                val = row[j] if j < len(row) else ""
                parts.append(f"<td>{_escape(str(val))}</td>")
            parts.append("</tr>")
        parts.append("</tbody></table>")

    parts.append("<h2>Agregats MACTOR</h2>")
    if mactor_result:
        for label, blob in (
            ("Mobilització actors", mactor_result.mobilisation_actors),
            ("Mobilització objectius", mactor_result.mobilisation_objectives),
            ("Matriu de convergències", mactor_result.convergences_matrix),
        ):
            if blob is not None:
                parts.append(f"<h3>{_escape(label)}</h3><pre class='json'>{_escape(_json_block(blob))}</pre>")
    else:
        parts.append("<p class='muted'>Sense resultats MACTOR calculats.</p>")

    parts.append("<h1>Anàlisi morfològic</h1>")
    if components:
        for c in components:
            parts.append(f"<h2>{_escape(c.code)} · {_escape(c.name)}</h2>")
            confs = c.configurations or []
            if isinstance(confs, list) and confs:
                parts.append('<table class="grid"><thead><tr><th>Etiqueta</th><th>Descripció</th></tr></thead><tbody>')
                for cfg in confs:
                    if isinstance(cfg, dict):
                        parts.append(
                            f"<tr><td>{_escape(cfg.get('label', ''))}</td>"
                            f"<td>{_escape(cfg.get('desc', ''))}</td></tr>"
                        )
                    else:
                        parts.append(f"<tr><td colspan='2'>{_escape(str(cfg))}</td></tr>")
                parts.append("</tbody></table>")
            else:
                parts.append("<p class='muted'>Sense configuracions.</p>")
    else:
        parts.append("<p class='muted'>Sense components morfològics.</p>")

    incompat = bundle.get("incompatibilities") or []
    if incompat:
        parts.append("<h2>Incompatibilitats morfològiques</h2>")
        parts.append('<table class="grid"><thead><tr><th>Component A</th><th>Config A</th><th>Component B</th><th>Config B</th></tr></thead><tbody>')
        for row in incompat[:30]:
            parts.append(
                f"<tr><td>{_escape(row.get('component_a'))}</td><td>{_escape(row.get('config_a'))}</td>"
                f"<td>{_escape(row.get('component_b'))}</td><td>{_escape(row.get('config_b'))}</td></tr>"
            )
        parts.append("</tbody></table>")

    smic = bundle.get("smic") or {}
    if smic and smic.get("initial_probs"):
        parts.append("<h2>SMIC · Probabilitats d'escenari</h2>")
        parts.append(f"<pre class='json'>{_escape(_json_block(smic))}</pre>")

    parts.append("<h1>Escenaris</h1>")
    if scenarios_list:
        for s in scenarios_list:
            parts.append(f"<h2>{_escape(s.name)} ({_escape(s.scenario_type or '')})</h2>")
            poss = getattr(s, "possibility", None) or "PLAUSIBLE"
            parts.append(
                f"<p class='muted'>Possibilitat (Zwicky): {_escape(poss)} · "
                f"Probabilitat (SMIC/OSINT): {_escape(s.probability)}</p>"
            )
            poss_rationale = getattr(s, "possibility_rationale", None) or ""
            if poss_rationale:
                parts.append(f"<p><strong>Possibilitat</strong> {_escape(poss_rationale)}</p>")
            if s.morphological_config:
                parts.append(f"<p><strong>Configuració</strong> {_escape(s.morphological_config)}</p>")
            nar = _escape(s.narrative or "").replace("\n", "<br/>")
            parts.append(f"<div>{nar}</div>")
    else:
        parts.append("<p class='muted'>Sense narratives d'escenari generades.</p>")

    _append_decision_annex_html(parts, bundle)
    _append_qual_quant_html(parts, bundle)

    return "<html><meta charset='utf-8'><body>" + "".join(parts) + "</body></html>"


def _write_pdf_sync(path: Path, bundle: dict[str, Any]) -> None:
    html_str = _html_report(bundle)
    try:
        render_pdf_from_html(html_str, path, base_url=str(path.parent.resolve()))
    except ExportBackendError as exc:
        raise RuntimeError(str(exc)) from exc


def _add_doc_cover(doc: Any, bundle: dict[str, Any]) -> None:
    """
    DOCX cover: iterate (text, size, grey_rgb) tuples only — add paragraph per row.
    """
    project: ProspectiveProject = bundle["project"]
    s = bundle["strings"]
    _, WD_ALIGN_PARAGRAPH, Pt, RGBColor, PRIMARY_RGB = _load_docx()
    subtitle_grey = RGBColor(0x33, 0x33, 0x33)
    meta_grey = RGBColor(0x66, 0x66, 0x66)
    accent_orange = RGBColor(0xFF, 0x6B, 0x35)

    cover_lines: list[tuple[str, Pt, RGBColor]] = [
        (s.report_title, Pt(26), PRIMARY_RGB),
        (project.title, Pt(22), subtitle_grey),
        (s.report_subtitle, Pt(13), subtitle_grey),
        (f"{s.project_id} {project.id}", Pt(11), meta_grey),
        (f"{s.generated}: {_utc_now_iso()}", Pt(11), meta_grey),
        ("______________________________________________", Pt(9), accent_orange),
    ]

    for txt, sz, grey in cover_lines:
        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run(txt)
        run.font.size = sz
        run.font.color.rgb = grey

    doc.add_page_break()


def _populate_doc_body(doc: Any, bundle: dict[str, Any]) -> None:
    project: ProspectiveProject = bundle["project"]
    vars_list: list = bundle["variables"]
    actors_list: list = bundle["actors"]
    objectives_list: list = bundle["objectives"]
    micmac: Optional[MICMACResult] = bundle["micmac"]
    mactor_result: Optional[MACTORResult] = bundle["mactor_result"]
    components: list = bundle["components"]
    scenarios_list: list = bundle["scenarios"]
    actor_codes = bundle["actor_codes"]
    objective_codes = bundle["objective_codes"]
    postures_matrix = bundle["postures_matrix"]

    _, _, _, _, PRIMARY_RGB = _load_docx()
    s = bundle["strings"]

    _append_executive_summary_docx(doc, bundle)

    h = doc.add_heading(s.project_section, level=1)
    h.runs[0].font.color.rgb = PRIMARY_RGB
    doc.add_paragraph(f"{s.hypothesis}: {project.hypothesis or ''}")
    doc.add_heading(s.context, level=2)
    doc.add_paragraph(project.context or "")

    _append_case_briefing_docx(doc, bundle)
    _append_osint_sources_docx(doc, bundle)
    _append_extraction_docx(doc, bundle)
    _append_retrospective_docx(doc, bundle)
    _append_actor_impact_docx(doc, bundle)
    _append_investment_docx(doc, bundle)

    _append_variable_profiles_docx(doc, bundle)

    doc.add_heading("Resultats MIC-MAC", level=1)
    if micmac:
        sector_rows = _format_sectors_table(micmac.sectors)
        if sector_rows:
            doc.add_heading("Classificació per sectors (Godet)", level=2)
            _doc_table(doc, ["Codi", "Sector", "Motricitat", "Dependència"], sector_rows)
        _append_micmac_reasoning_docx(doc, bundle)
        micmac_sections = (
            ("Sectors", micmac.sectors),
            ("Motricitat directa", micmac.motricite_direct),
            ("Dependència directa", micmac.dependence_direct),
            ("Matriu directa", micmac.matrix_direct),
            ("Matriu indirecta", micmac.matrix_indirect),
        )
        for title, blob in micmac_sections:
            if blob is not None:
                doc.add_heading(title, level=2)
                doc.add_paragraph(_json_block(blob))
    else:
        doc.add_paragraph("Sense MIC-MAC desat.")

    doc.add_heading("MACTOR", level=1)
    suggested_actors = bundle.get("suggested_actors") or []
    if suggested_actors:
        doc.add_heading("Actors suggerits per l'extracció OSINT", level=2)
        _doc_table(
            doc,
            ["Codi", "Nom", "Força", "Declaracions"],
            [[sa.get("code", ""), sa.get("name", ""), str(sa.get("force", "")), str(sa.get("statement_count", sa.get("n_statements", "")))] for sa in suggested_actors[:15]],
        )
    if actors_list:
        t = doc.add_table(rows=1, cols=4)
        hr = t.rows[0].cells
        hr[0].text = "Codi"
        hr[1].text = "Nom"
        hr[2].text = "Força"
        hr[3].text = "Fins"
        for a in actors_list:
            r = t.add_row().cells
            r[0].text = str(a.code)
            r[1].text = str(a.name)
            r[2].text = str(a.force_score)
            r[3].text = ", ".join(a.strategic_goals or [])
    if objectives_list:
        doc.add_heading("Objectius", level=2)
        ot = doc.add_table(rows=1, cols=2)
        ohr = ot.rows[0].cells
        ohr[0].text = "Codi"
        ohr[1].text = "Nom"
        for o in objectives_list:
            oc = ot.add_row().cells
            oc[0].text = str(o.code)
            oc[1].text = str(o.name)

    if actor_codes and objective_codes:
        doc.add_heading("Matriu de postures", level=2)
        cols = 1 + len(objective_codes)
        pt = doc.add_table(rows=1, cols=cols)
        head = pt.rows[0].cells
        head[0].text = "Actor \\ Objectiu"
        for j, oc in enumerate(objective_codes):
            head[j + 1].text = str(oc)
        for i, ac in enumerate(actor_codes):
            prow = pt.add_row().cells
            prow[0].text = str(ac)
            row_vals = postures_matrix[i] if i < len(postures_matrix) else []
            for j in range(len(objective_codes)):
                prow[j + 1].text = str(row_vals[j] if j < len(row_vals) else "")

    if mactor_result:
        doc.add_heading("Agregats MACTOR", level=2)
        for lbl, fld in (
            ("Mobilització actors", mactor_result.mobilisation_actors),
            ("Mobilització objectius", mactor_result.mobilisation_objectives),
            ("Convergències", mactor_result.convergences_matrix),
        ):
            if fld is not None:
                doc.add_paragraph(lbl + ":")
                doc.add_paragraph(_json_block(fld))

    doc.add_heading("Anàlisi morfològic", level=1)
    if components:
        for c in components:
            doc.add_heading(f"{c.code} — {c.name}", level=2)
            cfgs = c.configurations or []
            if isinstance(cfgs, list):
                ct = doc.add_table(rows=1, cols=2)
                chr_ = ct.rows[0].cells
                chr_[0].text = "Etiqueta"
                chr_[1].text = "Descripció"
                for cfg in cfgs:
                    if isinstance(cfg, dict):
                        cr = ct.add_row().cells
                        cr[0].text = str(cfg.get("label", ""))
                        cr[1].text = str(cfg.get("desc", ""))
                    else:
                        cr = ct.add_row().cells
                        cr[0].text = str(cfg)
                        cr[1].text = ""
    else:
        doc.add_paragraph("Sense components.")

    incompat = bundle.get("incompatibilities") or []
    if incompat:
        doc.add_heading("Incompatibilitats morfològiques", level=2)
        _doc_table(
            doc,
            ["Component A", "Config A", "Component B", "Config B"],
            [[row.get("component_a", ""), row.get("config_a", ""), row.get("component_b", ""), row.get("config_b", "")] for row in incompat[:30]],
        )
    smic = bundle.get("smic") or {}
    if smic and smic.get("initial_probs"):
        doc.add_heading("SMIC · Probabilitats d'escenari", level=2)
        doc.add_paragraph(_json_block(smic))

    doc.add_heading("Escenaris", level=1)
    if scenarios_list:
        for s in scenarios_list:
            doc.add_heading(s.name, level=2)
            poss = getattr(s, "possibility", None) or "PLAUSIBLE"
            doc.add_paragraph(
                f"Tipus: {s.scenario_type or '—'}  · "
                f"Possibilitat: {poss}  · Probabilitat: {s.probability or '—'}"
            )
            poss_rationale = getattr(s, "possibility_rationale", None) or ""
            if poss_rationale:
                doc.add_paragraph(f"Possibilitat: {poss_rationale}")
            if s.morphological_config:
                doc.add_paragraph(str(s.morphological_config))
            doc.add_paragraph(s.narrative or "")
    else:
        doc.add_paragraph("Sense escenaris generats.")

    _append_decision_annex_docx(doc, bundle)
    _append_qual_quant_docx(doc, bundle)


def _write_html_sync(bundle: dict[str, Any]) -> str:
    return _html_report(bundle)


async def export_html(
    db: AsyncSession,
    project_id: int,
    *,
    lang: str = "ca",
    include_decision_annex: bool = False,
    output_dir: str | Path | None = None,
    filename: str | None = None,
) -> dict[str, Any]:
    """Export full report as standalone HTML (print to PDF from browser on Windows)."""
    bundle = await _load_project_bundle(db, project_id)
    if not bundle:
        raise ValueError(f"Prospective project {project_id} not found")
    bundle = _prepare_export_bundle(bundle, lang)
    bundle = await _attach_decision_annex(db, bundle, include_decision_annex=include_decision_annex)

    directory = _ensure_reports_dir(output_dir)
    fname = filename or f"prospective_project_{project_id}.html"
    path = directory / fname
    html_str = await asyncio.to_thread(_write_html_sync, bundle)
    await asyncio.to_thread(path.write_text, html_str, encoding="utf-8")
    logger.info("Wrote prospective HTML export to %s", path)
    return {
        "file_path": str(path.resolve()),
        "format": "html",
        "project_id": project_id,
        "lang": bundle["lang"],
    }


def _write_docx_sync(bundle: dict[str, Any]) -> bytes:
    Document, _, _, _, _ = _load_docx()
    doc = Document()
    _add_doc_cover(doc, bundle)
    _populate_doc_body(doc, bundle)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


async def export_pdf(
    db: AsyncSession,
    project_id: int,
    *,
    lang: str = "ca",
    include_decision_annex: bool = False,
    output_dir: str | Path | None = None,
    filename: str | None = None,
) -> dict[str, Any]:
    """
    Export prospective project bundle to PDF. Returns metadata including absolute file_path.
    """
    bundle = await _load_project_bundle(db, project_id)
    if not bundle:
        raise ValueError(f"Prospective project {project_id} not found")
    bundle = _prepare_export_bundle(bundle, lang)
    bundle = await _attach_decision_annex(db, bundle, include_decision_annex=include_decision_annex)

    directory = _ensure_reports_dir(output_dir)
    fname = filename or f"prospective_project_{project_id}.pdf"
    path = directory / fname

    await asyncio.to_thread(_write_pdf_sync, path, bundle)
    logger.info("Wrote prospective PDF export to %s", path)

    return {
        "file_path": str(path.resolve()),
        "format": "pdf",
        "project_id": project_id,
        "lang": bundle["lang"],
    }


async def export_docx(
    db: AsyncSession,
    project_id: int,
    *,
    lang: str = "ca",
    include_decision_annex: bool = False,
    output_dir: str | Path | None = None,
    filename: str | None = None,
) -> dict[str, Any]:
    """
    Export prospective project bundle to DOCX. Returns metadata including absolute file_path.
    """
    bundle = await _load_project_bundle(db, project_id)
    if not bundle:
        raise ValueError(f"Prospective project {project_id} not found")
    bundle = _prepare_export_bundle(bundle, lang)
    bundle = await _attach_decision_annex(db, bundle, include_decision_annex=include_decision_annex)

    directory = _ensure_reports_dir(output_dir)
    fname = filename or f"prospective_project_{project_id}.docx"
    path = directory / fname

    data = await asyncio.to_thread(_write_docx_sync, bundle)
    await asyncio.to_thread(path.write_bytes, data)
    logger.info("Wrote prospective DOCX export to %s", path)

    return {
        "file_path": str(path.resolve()),
        "format": "docx",
        "project_id": project_id,
        "lang": bundle["lang"],
    }
